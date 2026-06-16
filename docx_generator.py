"""
docx_generator.py — Gera o relatório mensal de disponibilidade em DOCX (Word).

Reproduz o modelo do PDF em formato Word (python-docx), com:
  - Capa com KPIs consolidados
  - Metodologia e quadro consolidado
  - Seção por unidade (médias + Histórico de alertas + Resultados)

Menos formatação visual que o PDF (Word é mais limitado), mas totalmente editável
e compatível com Office/LibreOffice.
"""
from __future__ import annotations
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from charts import render_unit_graph_png

MESES_PT = ["", "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
            "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
INK = "0E2436"
INK_SOFT = "5A6B7B"
LINE = "DCE3EA"
SURFACE = "F7FAFB"
SURFACE_ALT = "F2F6F8"
TEAL = "0E7C86"
OK = "128A5E"
WARN = "D98A26"
CRIT = "C9384A"


def _pct(x):
    return f"{x:.2f}%".replace(".", ",")


def _fmt(n):
    return f"{int(n):,}".replace(",", ".")


def _minToHuman(m):
    h = m // 60
    mm = m % 60
    return f"{h}h {mm}min"


def _status_text(sla_pct):
    if sla_pct >= 99:
        return "Dentro do SLA"
    if sla_pct >= 90:
        return "Atenção"
    return "Crítico"


def _status_color(sla_pct):
    if sla_pct >= 99:
        return OK
    if sla_pct >= 90:
        return WARN
    return CRIT


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str = LINE, size: str = "8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _style_run(run, *, size=10, bold=False, color=INK, font_name="Arial", italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = font_name


def _set_cell_text(cell, lines: list[tuple[str, dict]], align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    for idx, (text, opts) in enumerate(lines):
        run = p.add_run(text)
        _style_run(run, **opts)
        if idx != len(lines) - 1:
            run.add_break()


def _style_table(table, first_col_width=None):
    table.style = "Table Grid"
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
    if first_col_width:
        for row in table.rows:
            row.cells[0].width = first_col_width


def _add_brand_block(doc: Document):
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = brand.add_run("K3G")
    _style_run(run, size=11, bold=True, color="FFFFFF")
    _set_cell_shading(doc.add_table(rows=1, cols=1).cell(0, 0), INK)


def _add_cover_brand(doc: Document):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(0.7)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, INK)
    _set_cell_border(cell, INK, "0")
    _set_cell_text(cell, [("K3G", {"size": 11, "bold": True, "color": "FFFFFF"})], align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_info_box(doc: Document, title: str, body: list[str]):
    table = doc.add_table(rows=1 + len(body), cols=1)
    table.autofit = True
    _style_table(table)
    _set_cell_shading(table.cell(0, 0), SURFACE_ALT)
    _set_cell_text(table.cell(0, 0), [(title, {"size": 11, "bold": True, "color": INK})])
    for idx, line in enumerate(body, start=1):
        _set_cell_shading(table.cell(idx, 0), SURFACE)
        _set_cell_text(table.cell(idx, 0), [(line, {"size": 10, "color": INK_SOFT})])


def _add_kpi_cards(doc: Document, items: list[tuple[str, str, str]]):
    table = doc.add_table(rows=1, cols=len(items))
    table.autofit = True
    for idx, (label, value, color) in enumerate(items):
        cell = table.rows[0].cells[idx]
        _set_cell_shading(cell, SURFACE)
        _set_cell_border(cell)
        _set_cell_text(
            cell,
            [
                (label.upper(), {"size": 8, "bold": True, "color": INK_SOFT}),
                (value, {"size": 18, "bold": True, "color": color}),
            ],
        )


def _style_header_row(row):
    for cell in row.cells:
        _set_cell_shading(cell, SURFACE_ALT)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                _style_run(run, size=9, bold=True, color=INK)


def _add_section_title(doc: Document, eyebrow: str, title: str):
    eye = doc.add_paragraph()
    eye.paragraph_format.space_after = Pt(2)
    run = eye.add_run(eyebrow.upper())
    _style_run(run, size=8, bold=True, color=TEAL)
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(title)
    _style_run(run, size=16, bold=True, color=INK)


def _add_key_value_table(doc: Document, rows: list[tuple[str, str]], first_col_fill=SURFACE):
    table = doc.add_table(rows=0, cols=2)
    _style_table(table, first_col_width=Inches(3.8))
    for key, value in rows:
        row = table.add_row()
        _set_cell_shading(row.cells[0], first_col_fill)
        _set_cell_text(row.cells[0], [(key, {"size": 10, "color": INK_SOFT})])
        _set_cell_text(row.cells[1], [(value, {"size": 10, "bold": True, "color": INK})], align=WD_ALIGN_PARAGRAPH.RIGHT)
    return table


def build_docx(report: dict, out_path: str) -> str:
    doc = Document()
    
    # Margem e estilo
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    p = report["periodo"]
    cons = report["consolidado"]
    units = report["unidades"]
    groups = report.get("grupos", [])
    selected_clients = report.get("clientes_selecionados", [])
    mes_nome = MESES_PT[p["mes"]]
    total = p["total_min"]
    selected_names = ", ".join(client["name"] for client in selected_clients) if selected_clients else "Filtro geral"

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)

    # ==================== Capa ====================
    _add_cover_brand(doc)
    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _style_run(eyebrow.add_run("RELATÓRIO DE DISPONIBILIDADE"), size=9, bold=True, color=TEAL)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("CLIENTES SELECIONADOS")
    _style_run(r, size=22, bold=True, color=INK)

    subtitle = doc.add_paragraph(f"Clientes Selecionados — {mes_nome} de {p['ano']}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_r = subtitle.runs[0]
    _style_run(subtitle_r, size=16, bold=True, color=INK)

    for text in (
        "Relatório multi-cliente de disponibilidade e SLA",
        "Fonte: monitoramento ICMP do Zabbix",
        f"Clientes selecionados: {selected_names}",
        f"Modo do relatório: {'Agrupado por cliente' if groups else 'Unificado'}",
    ):
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _style_run(p_info.add_run(text), size=10, color=INK_SOFT)

    doc.add_paragraph()  # espaço

    _add_kpi_cards(doc, [
        ("SLA médio consolidado", f"{cons['sla_medio_pct']:.2f}%".replace(".", ","), _status_color(cons["sla_medio_pct"])),
        ("Indisponibilidade total", f"{_fmt(cons['indisp_total_min'])} min", INK),
        ("Dentro do SLA", f"{cons['ok']}/{cons['unidades']}", OK),
        ("Crítico", f"{cons['crit']}", CRIT),
    ])

    doc.add_paragraph()

    # Metodologia
    _add_section_title(doc, "Metodologia", "Critérios de apuração")
    _add_info_box(doc, "Como o cálculo é feito", [
        f"A disponibilidade foi apurada com base na média das respostas de monitoramento ICMP (estado UP/DOWN) coletadas pelo Zabbix em intervalos regulares de 60 segundos, ao longo de {_fmt(total)} minutos ({total // 1440} dias).",
        "Fórmula: D = (To − Ti) / To × 100        Ti = To × (1 − D)",
        "To: período total de operação (min) · Ti: somatório das interrupções e intervalos com taxa de erro elevada · D: disponibilidade (decimal).",
    ])

    doc.add_paragraph()

    # Tabela consolidada
    _add_section_title(doc, "Consolidado", "Quadro consolidado das unidades")
    table = doc.add_table(rows=1, cols=5)
    _style_table(table, first_col_width=Inches(3.2))
    hdr_cells = table.rows[0].cells
    headers = ["Unidade", "SLA", "Indisp. (min)", "Latência", "Perda"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    _style_header_row(table.rows[0])

    for u in sorted(units, key=lambda x: x["sla_pct"]):
        row_cells = table.add_row().cells
        _set_cell_text(row_cells[0], [(u["nome"], {"size": 10, "bold": True, "color": INK})])
        _set_cell_text(row_cells[1], [(_pct(u["sla_pct"]), {"size": 10, "bold": True, "color": _status_color(u["sla_pct"])})], align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(row_cells[2], [(_fmt(u["downtime_min"]), {"size": 10, "color": INK})], align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(row_cells[3], [(f"{u['latency_ms']:.2f} ms", {"size": 10, "color": INK})], align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(row_cells[4], [(f"{u['packet_loss_pct']:.4f}%", {"size": 10, "color": INK})], align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph()

    if groups:
        _add_section_title(doc, "Clientes", "Quadros por cliente")
        for group in groups:
            client = group["cliente"]
            client_units = group["unidades"]
            client_cons = group["consolidado"]
            _add_section_title(doc, "Cliente", client["name"])
            _add_kpi_cards(doc, [
                ("SLA médio", _pct(client_cons["sla_medio_pct"]), _status_color(client_cons["sla_medio_pct"])),
                ("Indisponibilidade", f"{_fmt(client_cons['indisp_total_min'])} min", INK),
                ("Unidades", f"{client_cons['unidades']}", INK),
                ("Crítico", f"{client_cons['crit']}", CRIT),
            ])
            doc.add_paragraph()
            if not client_units:
                msg = doc.add_paragraph()
                _style_run(msg.add_run("Nenhuma unidade com itens ICMP válidos no período."), size=10, color=INK_SOFT, italic=True)
                continue
            client_table = doc.add_table(rows=1, cols=5)
            _style_table(client_table, first_col_width=Inches(3.2))
            hdr = client_table.rows[0].cells
            for i, header in enumerate(headers):
                hdr[i].text = header
            _style_header_row(client_table.rows[0])
            for u in sorted(client_units, key=lambda x: x["sla_pct"]):
                row = client_table.add_row().cells
                _set_cell_text(row[0], [(u["nome"], {"size": 10, "bold": True, "color": INK})])
                _set_cell_text(row[1], [(_pct(u["sla_pct"]), {"size": 10, "bold": True, "color": _status_color(u["sla_pct"])})], align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell_text(row[2], [(_fmt(u["downtime_min"]), {"size": 10, "color": INK})], align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell_text(row[3], [(f"{u['latency_ms']:.2f} ms", {"size": 10, "color": INK})], align=WD_ALIGN_PARAGRAPH.RIGHT)
                _set_cell_text(row[4], [(f"{u['packet_loss_pct']:.4f}%", {"size": 10, "color": INK})], align=WD_ALIGN_PARAGRAPH.RIGHT)
            doc.add_paragraph()

    # ==================== Seção por unidade ====================
    detail_units = []
    if groups:
        for group in groups:
            detail_units.extend(group["unidades"])
    else:
        detail_units = units

    for u in detail_units:
        doc.add_page_break()
        _add_section_title(doc, u.get("cliente", "Cliente"), u["nome"])
        addr = doc.add_paragraph()
        _style_run(addr.add_run(f"{u['local']} · IP {u['ip']}"), size=10, color=INK_SOFT)

        # Médias
        means_table = doc.add_table(rows=3, cols=2)
        _style_table(means_table, first_col_width=Inches(3.4))
        means_data = [
            ("Disponibilidade (SLA)", _pct(u["sla_pct"])),
            ("Latência Média", f"{u['latency_ms']:.2f} ms"),
            ("Perda de Pacotes", f"{u['packet_loss_pct']:.4f}%"),
        ]
        for i, (label, value) in enumerate(means_data):
            _set_cell_shading(means_table.rows[i].cells[0], SURFACE)
            _set_cell_text(means_table.rows[i].cells[0], [(label, {"size": 10, "color": INK_SOFT})])
            _set_cell_text(
                means_table.rows[i].cells[1],
                [(value, {"size": 12, "bold": True, "color": _status_color(u["sla_pct"]) if i == 0 else INK})],
                align=WD_ALIGN_PARAGRAPH.RIGHT,
            )

        doc.add_paragraph()

        graph_png = render_unit_graph_png(u.get("graphs"), u["nome"])
        if graph_png:
            _add_section_title(doc, "Gráficos", "Gráficos do Zabbix")
            doc.add_picture(BytesIO(graph_png), width=Inches(6.8))
            doc.add_paragraph()

        # Histórico de alertas
        _add_section_title(doc, "Incidentes", "Histórico de alertas")
        if u.get("incidentes"):
            inc_table = doc.add_table(rows=1, cols=4)
            _style_table(inc_table)
            hdr = inc_table.rows[0].cells
            for i, h in enumerate(["Início", "Resolvido", "Problema", "Duração"]):
                hdr[i].text = h
            _style_header_row(inc_table.rows[0])
            for inc in u["incidentes"]:
                r = inc_table.add_row().cells
                _set_cell_text(r[0], [(inc.get("inicio", ""), {"size": 9, "color": INK})])
                _set_cell_text(r[1], [(inc.get("fim", ""), {"size": 9, "color": INK})])
                _set_cell_text(r[2], [(inc.get("problema", ""), {"size": 9, "color": INK})])
                _set_cell_text(r[3], [(inc.get("duracao", ""), {"size": 9, "color": INK})], align=WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            msg = doc.add_paragraph()
            _style_run(msg.add_run("Sem registro de incidentes no período."), size=10, color=INK_SOFT, italic=True)

        doc.add_paragraph()

        # Resultados
        _add_section_title(doc, "Resultados", "Resultados")
        for text, color in (
            (f"Média observada: {u['availability']:.4f}", INK),
            (f"Tempo do mês (To): {_fmt(u['total_min'])} min", INK),
            (f"Tempo estimado de indisponibilidade (Ti): {_fmt(u['downtime_min'])} min", INK),
            (f"Tempo de disponibilidade: {_fmt(u['uptime_min'])} min", INK),
            (f"Disponibilidade SLA: {_pct(u['sla_pct'])} ({_status_text(u['sla_pct'])})", _status_color(u["sla_pct"])),
        ):
            p_item = doc.add_paragraph(style="List Bullet")
            _style_run(p_item.add_run(text), size=10, bold=text.startswith("Disponibilidade SLA"), color=color)

        doc.add_paragraph()

        # Resumo apurado
        _add_section_title(doc, "Resumo", "Resumo dos dados apurados")
        sum_data = [
            ("Tempo de indisponibilidade", f"{_fmt(u['downtime_min'])} minutos"),
            ("Tempo de disponibilidade", f"{_fmt(u['uptime_min'])} minutos"),
            ("Disponibilidade do serviço (SLA)", _pct(u["sla_pct"])),
            ("Latência média", f"{u['latency_ms']:.2f} ms"),
            ("Perda de pacotes (média)", f"{u['packet_loss_pct']:.4f}%"),
        ]
        _add_key_value_table(doc, sum_data)

    # Rodapé
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = footer.add_run(
        f"K3G Solutions LTDA · NOC 24×7 · Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
        f"Fonte: Zabbix · Cálculo SLA: D = (To − Ti) / To"
    )
    _style_run(footer_text, size=8, color=INK_SOFT, italic=True)

    doc.save(out_path)
    return out_path


# ----- Função para demonstração -----
def sample_report() -> dict:
    """Cria um relatório de exemplo para teste offline."""
    from report_generator import sample_report as pdf_sample
    return pdf_sample()


if __name__ == "__main__":
    build_docx(sample_report(), "Relatorio_Disponibilidade_PRF-AM_Marco-2026.docx")
    print("DOCX gerado.")
